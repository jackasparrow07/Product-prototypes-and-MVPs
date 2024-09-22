import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from scipy import stats
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
import io
import base64
import docx
from docx.shared import Inches
import xlsxwriter

st.set_page_config(page_title="Enhanced EDA Master", layout="wide", page_icon="📊")

# Custom CSS for improved design
st.markdown(
    """
    <style>
        .reportview-container {
            background: #f0f2f6;
        }
        .main .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        h1, h2, h3 {
            color: #1f2937;
        }
        .stButton>button {
            background-color: #3b82f6;
            color: white;
            border-radius: 0.375rem;
            padding: 0.5rem 1rem;
            border: none;
        }
        .stButton>button:hover {
            background-color: #2563eb;
        }
        .stTextInput>div>div>input {
            border-radius: 0.375rem;
        }
        .stSelectbox>div>div>select {
            border-radius: 0.375rem;
        }
    </style>
    """,
    unsafe_allow_html=True
)

@st.cache_data
def read_file(file):
    try:
        if file.name.endswith('.csv'):
            return pd.read_csv(file)
        elif file.name.endswith(('.xls', '.xlsx')):
            return pd.read_excel(file)
        else:
            st.error("Unsupported file format. Please upload a CSV or Excel file.")
            return None
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return None

def assess_data_quality(df):
    total_rows = len(df)
    missing_data = df.isnull().sum()
    missing_percentage = (missing_data / total_rows) * 100
    data_types = df.dtypes
    unique_values = df.nunique()

    quality_report = pd.DataFrame({
        'Data Type': data_types,
        'Missing Values': missing_data,
        'Missing Percentage': missing_percentage,
        'Unique Values': unique_values
    })

    return quality_report

def preprocess_data(df):
    st.subheader("Data Cleaning and Transformation")
    st.write("This section handles missing data and encodes categorical variables.")

    # Handle missing data
    for col in df.columns:
        if df[col].dtype in ['int64', 'float64']:
            df[col].fillna(df[col].median(), inplace=True)
            st.write(f"Filled missing values in '{col}' with median: {df[col].median():.2f}")
        elif df[col].dtype == 'object':
            df[col].fillna(df[col].mode()[0], inplace=True)
            st.write(f"Filled missing values in '{col}' with mode: {df[col].mode()[0]}")

    # Encode categorical variables
    cat_columns = df.select_dtypes(include=['object']).columns
    for col in cat_columns:
        df[col] = pd.Categorical(df[col]).codes
        st.write(f"Encoded categorical variable '{col}' with numerical codes.")

    return df

def generate_interactive_plots(df):
    st.subheader("Interactive Visualizations")
    st.write("Hover over the plots to see detailed information about each data point.")

    num_cols = df.select_dtypes(include=['int64', 'float64']).columns

    # Histogram for numerical columns
    for col in num_cols:
        fig = px.histogram(df, x=col, marginal="box", hover_data=df.columns)
        fig.update_layout(title=f"Distribution of {col}")
        st.plotly_chart(fig)
        st.markdown(f"""
        **How to interpret:**
        - The bars show the frequency of values in different ranges.
        - The box plot on top shows the median (middle line),
            1st and 3rd quartiles (box edges), and potential outliers (points).
        - Hover over bars to see exact counts and ranges.
        """)

    # Correlation heatmap
    corr_matrix = df[num_cols].corr()
    fig = px.imshow(corr_matrix, text_auto=True, aspect="auto")
    fig.update_layout(title="Correlation Heatmap")
    st.plotly_chart(fig)
    st.markdown("""
    **How to interpret:**
    - Colors represent the strength and direction of correlations between variables.
    - Red indicates positive correlation, blue indicates negative correlation.
    - Darker colors represent stronger correlations.
    - Hover over cells to see exact correlation values.
    """)

def extract_insights(df):
    st.subheader("Data Insights")

    # Basic statistics
    st.write("Basic Statistics:")
    st.write(df.describe())
    st.markdown("""
    **How to interpret:**
    - Count: Number of non-null values
    - Mean: Average value
    - Std: Standard deviation (measure of dispersion)
    - Min/Max: Minimum and maximum values
    - 25%/50%/75%: Quartile values (50% is the median)
    """)

    # Skewness and Kurtosis
    num_cols = df.select_dtypes(include=['int64', 'float64']).columns
    if not num_cols.empty:
        skewness = df[num_cols].apply(stats.skew)
        kurtosis = df[num_cols].apply(stats.kurtosis)

        st.write("Skewness:")
        st.write(skewness)
        st.markdown("""
        **How to interpret:**
        - Skewness measures the asymmetry of the distribution.
        - Positive skew: tail on right side (higher values)
        - Negative skew: tail on left side (lower values)
        - Values close to 0 indicate more symmetrical distribution
        """)

        st.write("Kurtosis:")
        st.write(kurtosis)
        st.markdown("""
        **How to interpret:**
        - Kurtosis measures the "tailedness" of the distribution.
        - Higher values indicate more extreme outliers.
        - Normal distribution has a kurtosis of 3.
        - Values > 3: heavier tails, < 3: lighter tails
        """)

    # Outlier detection
    st.subheader("Outlier Detection")
    for col in num_cols:
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)][col]
        if not outliers.empty:
            st.write(f"Outliers in {col}:")
            st.write(outliers)
            fig = px.box(df, y=col)
            st.plotly_chart(fig)
            st.markdown(f"""
            **How to interpret:**
            - Points beyond the whiskers are considered potential outliers.
            - Lower bound: {lower_bound:.2f}
            - Upper bound: {upper_bound:.2f}
            - Number of outliers: {len(outliers)}
            - Outliers may be errors or genuine extreme values.
            """)

def feature_importance(df):
    st.subheader("Feature Importance")

    num_cols = df.select_dtypes(include=['int64', 'float64']).columns
    if len(num_cols) < 2:
        st.write("Not enough numerical columns for feature importance analysis.")
        return

    X = df[num_cols].drop(num_cols[0], axis=1)
    y = df[num_cols[0]]

    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X, y)

    importance = pd.DataFrame({'feature': X.columns, 'importance': model.feature_importances_})
    importance = importance.sort_values('importance', ascending=False)

    fig = px.bar(importance, x='feature', y='importance', title="Feature Importance")
    st.plotly_chart(fig)

    st.markdown("""
    **How to interpret:**
    - Bars represent the relative importance of each feature in predicting the target variable.
    - Higher values indicate more important features.
    - This analysis helps identify which variables have the strongest relationship with the target.
    """)

def time_series_analysis(df):
    st.subheader("Time Series Analysis")

    date_cols = df.select_dtypes(include=['datetime64']).columns
    if len(date_cols) == 0:
        st.write("No datetime columns found for time series analysis.")
        return

    date_col = st.selectbox("Select date column for time series analysis:", date_cols)
    value_col = st.selectbox("Select value column for time series analysis:", df.select_dtypes(include=['int64', 'float64']).columns)

    df_ts = df[[date_col, value_col]].set_index(date_col)
    df_ts = df_ts.resample('D').mean()

    fig = px.line(df_ts, x=df_ts.index, y=value_col, title=f"Time Series of {value_col}")
    st.plotly_chart(fig)

    st.markdown("""
    **How to interpret:**
    - The line shows the trend of the selected value over time.
    - Look for patterns such as seasonality, trends, or sudden changes.
    - Hover over points to see exact values at specific dates.
    """)

def simple_ml_prediction(df):
    st.subheader("Simple Machine Learning Prediction")

    num_cols = df.select_dtypes(include=['int64', 'float64']).columns
    if len(num_cols) < 2:
        st.write("Not enough numerical columns for prediction.")
        return

    target = st.selectbox("Select target variable for prediction:", num_cols)
    features = [col for col in num_cols if col != target]

    X = df[features]
    y = df[target]

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)

    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X_train_scaled, y_train)

    train_score = model.score(X_train_scaled, y_train)
    test_score = model.score(X_test_scaled, y_test)

    st.write(f"Model R-squared score on training data: {train_score:.4f}")
    st.write(f"Model R-squared score on test data: {test_score:.4f}")

    st.markdown("""
    **How to interpret:**
    - R-squared score ranges from 0 to 1, with 1 being a perfect fit.
    - Higher scores indicate better predictive performance.
    - A large difference between training and test scores may indicate overfitting.
    """)

def save_as_markdown(content):
    return content

def save_as_docx(content):
    doc = docx.Document()
    doc.add_heading("EDA Report", 0)
    doc.add_paragraph(content)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def save_as_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def save_as_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Sheet1', index=False)
    return output.getvalue()

def main():
    st.title("📊 Enhanced EDA Master")
    st.subheader("Comprehensive Automatic EDA and Data Analysis")

    uploaded_file = st.file_uploader("Choose a CSV or Excel file", type=["csv", "xlsx", "xls"])

    if uploaded_file is not None:
        df = read_file(uploaded_file)

        if df is not None:
            st.subheader("Data Preview")
            st.write(df.head())

            st.subheader("Data Quality Assessment")
            quality_report = assess_data_quality(df)
            st.write(quality_report)

            if st.button("Preprocess Data"):
                df = preprocess_data(df)
                st.success("Data preprocessed successfully!")

            generate_interactive_plots(df)
            extract_insights(df)
            feature_importance(df)
            time_series_analysis(df)
            simple_ml_prediction(df)

            # Save options
            st.subheader("Save Results")
            save_format = st.multiselect("Select formats to save:", ["Markdown", "Word", "CSV", "Excel"])

            if st.button("Save Results"):
                if "Markdown" in save_format:
                    markdown_content = save_as_markdown(str(df.describe()))
                    st.download_button("Download Markdown", markdown_content, "eda_report.md")

                if "Word" in save_format:
                    docx_content = save_as_docx(str(df.describe()))
                    st.download_button("Download Word", docx_content, "eda_report.docx")

                if "CSV" in save_format:
                    csv_content = save_as_csv(df)
                    st.download_button("Download CSV", csv_content, "eda_data.csv")

                if "Excel" in save_format:
                    excel_content = save_as_excel(df)
                    st.download_button("Download Excel", excel_content, "eda_data.xlsx")

if __name__ == "__main__":
    main()